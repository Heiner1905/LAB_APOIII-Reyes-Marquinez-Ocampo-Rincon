"""
Lee los resultados del notebook tarea3 ejecutado y actualiza las dos tablas
del informe con los números resultantes (R²/RMSE/MAE para regresión,
Accuracy/F1/Kappa para clasificación). También extrae el AUC de la sección 16.
"""
import json
import re
from pathlib import Path
from docx import Document

ROOT = Path('/home/heiner/Documentos/apo_III/LAB_APOIII-Reyes-Marquinez-Ocampo-Rincon')
NB_T3 = ROOT / 'tarea3_modelamiento.ipynb'
DOCX = ROOT / 'Informe APO Lab.docx'


def get_text_outputs(cell):
    out = []
    for o in cell.get('outputs', []):
        if o.get('output_type') == 'stream':
            txt = o.get('text', '')
            if isinstance(txt, list):
                txt = ''.join(txt)
            out.append(txt)
        elif o.get('output_type') == 'execute_result':
            data = o.get('data', {})
            if 'text/plain' in data:
                txt = data['text/plain']
                if isinstance(txt, list):
                    txt = ''.join(txt)
                out.append(txt)
    return '\n'.join(out)


def parse_table_reg(text):
    """Espera formato:
       Modelo Target R² RMSE MAE
       OLS TCH 0.16 ...
    """
    rows = []
    for line in text.splitlines():
        line = line.strip()
        # Modelos: OLS / Ridge / Random Forest / XGBoost
        m = re.match(r'^(OLS|Ridge|Random Forest|XGBoost)\s+([%A-Za-zÁÉÍÓÚáéíóúÑñ.\s]+?)\s+(\d+\.\d+)\s+(\d+\.\d+)\s+(\d+\.\d+)\s*$', line)
        if m:
            rows.append({
                'Modelo': m.group(1),
                'Target': m.group(2).strip(),
                'R²':   m.group(3),
                'RMSE': m.group(4),
                'MAE':  m.group(5),
            })
    return rows


def parse_table_clf(text):
    """Espera líneas tipo: Modelo  Target  Accuracy  Precision  Recall  F1-macro  Kappa
    """
    rows = []
    for line in text.splitlines():
        line = line.strip()
        m = re.match(
            r'^(Logística L2|KNN|Random Forest|XGBoost)\s+(nivel TCH|nivel sacarosa)\s+'
            r'(\d+\.\d+)\s+(\d+\.\d+)\s+(\d+\.\d+)\s+(\d+\.\d+)\s+(\d+\.\d+)\s*$',
            line
        )
        if m:
            rows.append({
                'Modelo': m.group(1),
                'Target': m.group(2),
                'Accuracy':  m.group(3),
                'Precision': m.group(4),
                'Recall':    m.group(5),
                'F1-macro':  m.group(6),
                'Kappa':     m.group(7),
            })
    return rows


def main():
    nb = json.load(open(NB_T3))

    # Cell 19 imprime tabla_reg y cell 37 tabla_clf (índices se mantienen de T3 actual)
    text_reg = get_text_outputs(nb['cells'][19])
    text_clf = get_text_outputs(nb['cells'][37])

    rows_reg = parse_table_reg(text_reg)
    rows_clf = parse_table_clf(text_clf)

    if not rows_reg:
        print('WARN: no se pudo parsear tabla regresión, dump:')
        print(text_reg[:500])
    if not rows_clf:
        print('WARN: no se pudo parsear tabla clasificación, dump:')
        print(text_clf[:500])

    print(f'Filas regresión   : {len(rows_reg)}')
    print(f'Filas clasificación: {len(rows_clf)}')
    for r in rows_reg:
        print(' ', r)
    for r in rows_clf:
        print(' ', r)

    # Update docx tables
    doc = Document(DOCX)
    if len(doc.tables) < 2:
        print('ERROR: el informe no tiene 2 tablas, no se actualiza.')
        return

    # Tabla 1: regresión – orden esperado [OLS TCH, Ridge TCH, RF TCH, XGB TCH, OLS Sac, ...]
    tbl_reg = doc.tables[0]
    if rows_reg and len(tbl_reg.rows) - 1 == len(rows_reg):
        # Mantener orden por (Target, Modelo)
        order_reg = ['OLS', 'Ridge', 'Random Forest', 'XGBoost']
        rows_reg_sorted = sorted(
            rows_reg,
            key=lambda r: (0 if r['Target'] == 'TCH' else 1, order_reg.index(r['Modelo']))
        )
        for i, r in enumerate(rows_reg_sorted):
            row = tbl_reg.rows[i + 1]
            row.cells[0].text = r['Modelo']
            row.cells[1].text = r['Target']
            row.cells[2].text = r['R²']
            row.cells[3].text = r['RMSE']
            row.cells[4].text = r['MAE']
        print('Tabla 1 (regresión) actualizada.')
    else:
        print(f'SKIP tabla regresión: filas={len(rows_reg)}, esperadas={len(tbl_reg.rows)-1}')

    # Tabla 2: clasificación
    tbl_clf = doc.tables[1]
    if rows_clf and len(tbl_clf.rows) - 1 == len(rows_clf):
        order_clf = ['Logística L2', 'KNN', 'Random Forest', 'XGBoost']
        rows_clf_sorted = sorted(
            rows_clf,
            key=lambda r: (0 if r['Target'] == 'nivel TCH' else 1, order_clf.index(r['Modelo']))
        )
        for i, r in enumerate(rows_clf_sorted):
            row = tbl_clf.rows[i + 1]
            row.cells[0].text = r['Modelo']
            row.cells[1].text = r['Target']
            row.cells[2].text = r['Accuracy']
            row.cells[3].text = r['F1-macro']
            row.cells[4].text = r['Kappa']
        print('Tabla 2 (clasificación) actualizada.')
    else:
        print(f'SKIP tabla clasificación: filas={len(rows_clf)}, esperadas={len(tbl_clf.rows)-1}')

    doc.save(DOCX)
    print(f'\nInforme actualizado: {DOCX}')


if __name__ == '__main__':
    main()
