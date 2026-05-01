"""
Extrae las figuras clave de los notebooks ejecutados y las incrusta en el informe.
La consigna pide visualizaciones de apoyo (importancia, matrices de confusión, real vs predicho).

Ejecutar DESPUÉS de que tarea2 y tarea3 estén ejecutados completamente.
"""
import json
import base64
from pathlib import Path
from io import BytesIO
from docx import Document
from docx.shared import Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

ROOT = Path('/home/heiner/Documentos/apo_III/LAB_APOIII-Reyes-Marquinez-Ocampo-Rincon')
DOCX = ROOT / 'Informe APO Lab.docx'

# Mapeo de figuras a extraer: (notebook, índice de celda, descripción para cabecera)
# Los índices se leen del notebook ejecutado; cada figura es el primer output
# de tipo image/png en la celda indicada.
FIGURES_TO_EMBED = [
    # Tarea 2 — distribuciones y correlaciones
    ('tarea2_eda_preprocesamiento.ipynb', 18,
     'Figura 1. Distribución de TCH y %Sac.Caña – HISTORICO_SUERTES.'),
    ('tarea2_eda_preprocesamiento.ipynb', 27,
     'Figura 2. Matriz de correlación – variables numéricas (regresión).'),
    # Tarea 3 — diagnósticos y resultados
    ('tarea3_modelamiento.ipynb', 21,
     'Figura 3. Real vs Predicho – modelos de regresión.'),
    ('tarea3_modelamiento.ipynb', 22,
     'Figura 4. Feature Importance – Random Forest y XGBoost (regresión).'),
    ('tarea3_modelamiento.ipynb', 35,
     'Figura 5. Matrices de confusión – clasificación.'),
    ('tarea3_modelamiento.ipynb', 43,
     'Figura 6. Curvas ROC One-vs-Rest – clasificación.'),
    ('tarea3_modelamiento.ipynb', 45,
     'Figura 7. Curvas de aprendizaje – mejores modelos.'),
]


def extract_png_from_cell(nb_path, cell_idx):
    nb = json.load(open(nb_path))
    if cell_idx >= len(nb['cells']):
        return None
    cell = nb['cells'][cell_idx]
    for output in cell.get('outputs', []):
        data = output.get('data', {})
        if 'image/png' in data:
            png = data['image/png']
            if isinstance(png, list):
                png = ''.join(png)
            return base64.b64decode(png)
    return None


def main():
    doc = Document(DOCX)

    # Buscar al final del informe el lugar para añadir el anexo de figuras.
    # Añadimos un Anexo con las figuras (que NO cuenta para las 5 páginas, según consigna).
    doc.add_page_break()
    doc.add_paragraph('Anexo A – Figuras de Apoyo', style='Heading 1')
    doc.add_paragraph(
        'Las siguientes figuras complementan el cuerpo del informe y soportan los '
        'resultados discutidos. Todas son generadas por los notebooks ejecutados '
        '(tarea2_eda_preprocesamiento.ipynb y tarea3_modelamiento.ipynb).'
    )

    embedded = 0
    for nb_path, cell_idx, caption in FIGURES_TO_EMBED:
        full_path = ROOT / nb_path
        if not full_path.exists():
            print(f'  SKIP {nb_path}: not found')
            continue
        png_bytes = extract_png_from_cell(full_path, cell_idx)
        if png_bytes is None:
            print(f'  SKIP {nb_path} cell {cell_idx}: no image found')
            continue

        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(BytesIO(png_bytes), width=Cm(15))

        cap = doc.add_paragraph(caption)
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap_run = cap.runs[0]
        cap_run.italic = True
        cap_run.font.size = doc.styles['normal'].font.size

        embedded += 1
        print(f'  Embed: {nb_path} cell {cell_idx} → {caption[:50]}')

    doc.save(DOCX)
    print(f'\nFiguras incrustadas: {embedded}/{len(FIGURES_TO_EMBED)}')
    print(f'Informe actualizado: {DOCX}')


if __name__ == '__main__':
    main()
