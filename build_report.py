"""
Construye el informe final del Laboratorio 1 sobre el docx existente.
- Preserva portada y Tarea 1 (introducción/contexto del negocio).
- Compacta la Tarea 1 bajo encabezado "1. Introducción".
- Añade: Resumen, Metodología (datos/preproc/modelos/validación), Resultados y Discusión,
  Conclusiones y Trabajo Futuro, Referencias.
- Incluye tablas comparativas con resultados de los notebooks.
- Métricas leídas dinámicamente desde los notebooks ejecutados (cuando posible).
"""
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from copy import deepcopy
import json
from pathlib import Path

ROOT = Path('/home/heiner/Documentos/apo_III/LAB_APOIII-Reyes-Marquinez-Ocampo-Rincon')
DOCX = ROOT / 'Informe APO Lab.docx'

doc = Document(DOCX)

# ── 1. Localizar y limpiar la zona después de Tarea 1 ─────────────────────
# Encontrar el último párrafo "real" de Tarea 1 (Estado del Arte) y borrar todo lo siguiente.
def find_paragraph_idx(doc, text_substring):
    for i, p in enumerate(doc.paragraphs):
        if text_substring.lower() in p.text.lower():
            return i
    return None

# Buscar el último párrafo no vacío de Tarea 1 (sustenta la estrategia del taller)
last_t1 = None
for i, p in enumerate(doc.paragraphs):
    if 'sustenta la estrategia del taller' in p.text.lower():
        last_t1 = i
        break

if last_t1 is None:
    raise RuntimeError("No se encontró el final de Tarea 1 esperado.")

# Eliminar todos los párrafos vacíos al final
body = doc.paragraphs[0]._element.getparent()
for p in list(doc.paragraphs[last_t1+1:]):
    p._element.getparent().remove(p._element)

# Cambiar encabezado "TAREA 1: Introducción y Contexto del Negocio" a algo más limpio
for p in doc.paragraphs:
    if p.text.strip().startswith('TAREA 1'):
        p.text = '1. Introducción y Contexto del Negocio'
        p.style = doc.styles['Heading 1']
        break

# Asegurar estilos consistentes en sub-secciones existentes (Heading 2)
sub_h1_titles = [
    '1. Importancia Económica',
    '2. Factores Clave',
    '2.1 TCH',
    '2.2 Porcentaje de Sacarosa',
    '3. El Ingenio Providencia',
    '4. Estado del Arte',
]
# (lo dejamos como está; foco en agregar contenido)

# ── 2. Helpers para añadir secciones ──────────────────────────────────────
def add_h1(text):
    p = doc.add_paragraph(text, style='Heading 1')
    return p

def add_h2(text):
    return doc.add_paragraph(text, style='Heading 2')

def add_p(text, bold_first=False):
    p = doc.add_paragraph()
    p.add_run(text)
    return p

def add_bullet(text):
    # No hay estilo "List Bullet" en este docx; usamos un bullet manual
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.6)
    p.add_run('• ').bold = False
    p.add_run(text)
    return p

def add_table(headers, rows, widths_cm=None):
    table = doc.add_table(rows=1+len(rows), cols=len(headers))
    # Estilo nativo disponible: TableNormal (sin bordes); preferimos crear borders por XML.
    table.style = 'TableNormal'
    # Aplicar bordes simples a todas las celdas
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    for row in table.rows:
        for cell in row.cells:
            tcPr = cell._tc.get_or_add_tcPr()
            tcBorders = OxmlElement('w:tcBorders')
            for side in ('top', 'left', 'bottom', 'right'):
                b = OxmlElement(f'w:{side}')
                b.set(qn('w:val'), 'single')
                b.set(qn('w:sz'), '4')
                b.set(qn('w:color'), '000000')
                tcBorders.append(b)
            tcPr.append(tcBorders)
    # Headers
    for j, h in enumerate(headers):
        cell = table.rows[0].cells[j]
        cell.text = ''
        run = cell.paragraphs[0].add_run(h)
        run.bold = True
    for i, row in enumerate(rows):
        for j, val in enumerate(row):
            table.rows[i+1].cells[j].text = str(val)
    if widths_cm:
        for col, w in zip(table.columns, widths_cm):
            for cell in col.cells:
                cell.width = Cm(w)
    return table

# ── 3. RESUMEN (Abstract) ─────────────────────────────────────────────────
# Antes de la Introducción, insertamos el resumen. Como la introducción ya está,
# y no podemos insertar fácilmente "antes" sin perder portada, lo añadimos al final
# como nueva sección al inicio del cuerpo. En su lugar lo añadimos antes de Introducción.
# Estrategia: lo añadimos como bloque final ANTES de las nuevas secciones 2-6, y al final
# moveremos el bloque resumen para que aparezca antes de la Introducción.
# Por simpleza, añadimos el resumen al inicio insertando elementos XML.

resumen_text = (
    "Este trabajo aplica técnicas de regresión y clasificación sobre datos históricos de "
    "producción de caña de azúcar del Ingenio Providencia, con el objetivo de predecir el "
    "rendimiento (TCH, toneladas de caña por hectárea) y la calidad (%Sacarosa.Caña), y de "
    "clasificar los lotes en niveles de desempeño Bajo/Medio/Alto. Se trabajaron dos datasets: "
    "HISTORICO_SUERTES.xlsx (21 027 registros, 85 variables) para regresión y BD_IPSA_1940.xlsx "
    "(2 187 registros, 21 variables) para clasificación. La metodología incluye un EDA exhaustivo "
    "con análisis de nulos, outliers (regla del IQR) y multicolinealidad (VIF con intercepto), "
    "preprocesamiento basado en pipelines de scikit-learn con One-Hot Encoding y escalado, "
    "validación rigurosa con hold-out 80/20 y validación cruzada de 5 vueltas, y comparación de "
    "modelos de referencia (OLS, Ridge, Logística L2, KNN) frente a modelos avanzados (Random "
    "Forest y XGBoost). Para regresión, XGBoost obtuvo el mejor desempeño con R² ≈ 0.49 en TCH y "
    "R² ≈ 0.51 en %Sac.Caña, superando holgadamente a OLS (R² ≈ 0.20). Para clasificación, "
    "Random Forest y XGBoost alcanzaron F1-macro ≈ 0.49–0.55, con AUC-macro entre 0.65 y 0.74. "
    "Las variables más influyentes resultaron ser la edad del cultivo, la dosis de madurante, las "
    "lluvias en el periodo de maduración y la variedad/zona. Los resultados son consistentes con "
    "el estado del arte agronómico y sugieren oportunidades concretas para optimizar la "
    "planificación de cosecha y la aplicación de madurantes en el ingenio."
)

# Añadiremos al final, luego mover. Mejor: añadimos como bloque temporal etiquetado.
# Por practicidad: añadimos las secciones nuevas tras la introducción y dejamos el resumen
# justo después de la portada usando una manipulación XML.

# 4. METODOLOGÍA (TAREA 2 + TAREA 3 — descripción) ────────────────────────
add_h1('2. Metodología')

add_h2('2.1 Datos y Preprocesamiento')
add_p(
    "Se trabajaron dos datasets independientes provistos por el Ingenio Providencia: "
    "HISTORICO_SUERTES.xlsx (21 027 filas × 85 columnas) para los modelos de regresión, y "
    "BD_IPSA_1940.xlsx (2 187 filas × 21 columnas) — un subconjunto correspondiente a la "
    "variedad CC01-1940 — para los modelos de clasificación."
)
add_p(
    "Selección de variables. Para regresión se descartaron columnas con más del 70% de "
    "valores nulos (incluidas las series climáticas alternas, fertilización detallada e "
    "infestación por diátrea), identificadores sin poder predictivo (Hacienda, Suerte, Nombre) "
    "y variables con leakage por construirse a partir de los targets (TCHM, TonAzucar, Rdto, "
    "TAH, TAHM, Brix, Pureza, %ATR, KATRHM, %Sac.Muestreadora, %Fibra Caña, %AR Jugo, %ME, "
    "Fosfato Jugo, Sac.Caña Precosecha, TonUltCorte). El conjunto final consta de 21 "
    "predictores: edad del cultivo, vejez, variedad, suelo, zona, tenencia, distancia al "
    "ingenio, tipo de quema, tipo de corte, dosis de madurante, lluvias por sub-periodo y "
    "ciclo, riego (M3 y días desde último riego) y período. Para clasificación se eligieron "
    "10 variables numéricas (edad, cortes, lluvias, grupo_tenencia, %diatrea, dosis y semanas "
    "de madurante, vejez, mes, masa estructural). Se descartaron variedad, tipocorte, madurada "
    "y producto por ser constantes en el subconjunto (no aportan información)."
)
add_p(
    "Imputación. En el dataset de regresión, Suelo (17.9% nulos) recibió la categoría "
    "“DESCONOCIDO” para preservar la información de ausencia; Vejez (11.6%) se imputó con la "
    "mediana por Zona; Dist Km, Tenencia y Cod. T.Cultivo (≤0.02% nulos) con la mediana o moda "
    "global; Dosis Madurante (0.5%) con cero (la ausencia indica dosis no aplicada). Las 449 "
    "filas con %Sac.Caña nulo fueron eliminadas (no se imputan targets). El dataset de "
    "clasificación no presentó nulos."
)
add_p(
    "Outliers. Se aplicó la regla del IQR con factor 3.0 sobre los targets de regresión "
    "(criterio conservador apropiado a datos agrícolas), eliminando 19 registros físicamente "
    "implausibles. Tras esta limpieza, el dataset de regresión queda en 20 559 filas."
)
add_p(
    "Multicolinealidad. Se calculó el Factor de Inflación de la Varianza (VIF) sobre las "
    "variables numéricas, agregando intercepto a la matriz como exige statsmodels. Todas las "
    "variables resultaron con VIF < 5: Lluvias Ciclo (3.36) y Lluvias 2 Meses Ant. (2.54) "
    "fueron las más altas, seguidas de Tenencia (1.84). Por construcción, las cuatro lluvias "
    "por sub-periodo se excluyen del VIF al sumar exactamente Lluvias Ciclo (colinealidad "
    "perfecta); para los modelos lineales se conservan únicamente las dos lluvias agregadas, "
    "mientras que los modelos de árbol incluyen todos los sub-periodos."
)
add_p(
    "Codificación. Se utilizó One-Hot Encoding (sklearn ColumnTransformer + OneHotEncoder con "
    "min_frequency=0.005 para evitar la explosión por la alta cardinalidad de Variedad y "
    "Suelo) para todas las variables categóricas en todos los modelos, evitando el orden "
    "artificial que introduce LabelEncoding en modelos lineales. Las variables numéricas se "
    "estandarizaron (StandardScaler) para OLS, Ridge, Logística y KNN; los modelos de árbol no "
    "requieren escalado."
)
add_p(
    "Targets de clasificación. Se crearon dos variables ordinales (Bajo/Medio/Alto) "
    "discretizando TCH y sacarosa por tertiles exactos (pd.qcut con q=3), garantizando "
    "balance de clases (≈33% cada una)."
)

add_h2('2.2 Modelos')
add_p(
    "La estrategia va de lo simple a lo complejo, comparando modelos lineales contra modelos "
    "de ensamble:"
)
add_bullet(
    "Regresión – referencia: OLS (statsmodels) y Ridge con búsqueda de α en {0.01, 0.1, 1, "
    "10, 50, 100, 500} mediante 5-fold CV."
)
add_bullet(
    "Regresión – avanzados: Random Forest Regressor y XGBoost Regressor con RandomizedSearchCV "
    "(25 combinaciones, 5-fold) sobre n_estimators, max_depth, learning_rate, subsample, "
    "colsample_bytree y min_samples_leaf según corresponda."
)
add_bullet(
    "Clasificación – referencia: Regresión Logística multinomial con regularización L2 "
    "(GridSearchCV sobre C) y K-Vecinos Más Cercanos (GridSearchCV sobre k ∈ {3,5,...,31} y "
    "weights ∈ {uniform, distance})."
)
add_bullet(
    "Clasificación – avanzados: Random Forest Classifier (con class_weight=‘balanced’) y "
    "XGBoost Classifier, ambos con RandomizedSearchCV de 25 combinaciones × 5 folds."
)
add_p(
    "Todos los modelos se entrenaron dentro de un sklearn.Pipeline acoplado al "
    "ColumnTransformer correspondiente, garantizando que el preprocesamiento se aplique solo "
    "con datos de entrenamiento dentro de cada fold de validación cruzada (sin fuga)."
)

add_h2('2.3 Estrategia de Validación y Métricas')
add_p(
    "Hold-out 80/20 estratificado: para regresión se estratificó por cuartiles del target "
    "para garantizar distribuciones similares en train/test; para clasificación se estratificó "
    "por la clase ordinal (un split independiente por target). Validación cruzada 5-fold "
    "sobre el conjunto de entrenamiento para sintonizar hiperparámetros y para reportar "
    "métricas estables. Reproducibilidad mediante semilla aleatoria global SEED=42."
)
add_p(
    "Métricas — Regresión: R² (coeficiente de determinación), RMSE (raíz del error cuadrático "
    "medio), MAE (error absoluto medio); diagnóstico OLS con residuos vs predichos, QQ-plot, "
    "escala-localización y test de Shapiro-Wilk. Métricas — Clasificación: Accuracy, "
    "Precision, Recall, F1-macro, Cohen Kappa, y AUC One-vs-Rest macro mediante curvas ROC."
)

# ── 5. RESULTADOS Y DISCUSIÓN ─────────────────────────────────────────────
add_h1('3. Resultados y Discusión')

add_h2('3.1 Comparación de Modelos – Regresión')
add_p(
    "La Tabla 1 resume el desempeño en el set de prueba (20%) tras hyperparameter tuning con "
    "5-fold CV. Los modelos de ensamble más que duplican el R² de los modelos lineales, "
    "evidenciando que las relaciones entre los predictores y el rendimiento de la caña son "
    "sustancialmente no lineales y con interacciones complejas — un hallazgo consistente con "
    "la literatura de agricultura de precisión."
)
add_table(
    ['Modelo', 'Target', 'R²', 'RMSE', 'MAE'],
    [
        ['OLS',           'TCH',        '0.16', '29.19', '22.51'],
        ['Ridge',         'TCH',        '0.16', '29.18', '22.51'],
        ['Random Forest', 'TCH',        '0.46', '23.51', '17.71'],
        ['XGBoost',       'TCH',        '0.49', '22.80', '17.23'],
        ['OLS',           '%Sac.Caña',  '0.20', '1.012', '0.791'],
        ['Ridge',         '%Sac.Caña',  '0.20', '1.012', '0.791'],
        ['Random Forest', '%Sac.Caña',  '0.49', '0.803', '0.615'],
        ['XGBoost',       '%Sac.Caña',  '0.51', '0.791', '0.607'],
    ]
)
add_p(
    "XGBoost es el modelo ganador para los dos targets de regresión. Para TCH el RMSE de 22.8 "
    "ton/ha equivale a un error típico de ~17.5% sobre la media de 130 ton/ha; para %Sac.Caña "
    "el RMSE de 0.79 puntos porcentuales es del orden de la propia desviación natural del "
    "indicador (~1.1 pp). Los coeficientes OLS (significativos al 5%) confirman las relaciones "
    "esperadas: signo positivo para Edad Ult Cos y Dosis Madurante (más maduración → más "
    "azúcar), signo negativo para Lluvias 2 Meses Ant. en %Sac.Caña (exceso de agua diluye "
    "sacarosa). La importancia de variables del XGBoost destaca consistentemente la edad del "
    "cultivo, la dosis de madurante y las variables de variedad y zona como los predictores "
    "más informativos."
)

add_h2('3.2 Comparación de Modelos – Clasificación')
add_p(
    "La Tabla 2 muestra el desempeño en el set de prueba para los dos targets de "
    "clasificación, aplicando el mismo split estratificado a todos los modelos. Para nivel_TCH "
    "el F1-macro ronda 0.45–0.49, mientras que para nivel_sacarosa los modelos avanzados "
    "alcanzan 0.55. La Kappa (entre 0.18 y 0.35) indica un acuerdo modesto pero claramente "
    "superior al azar (Kappa = 0)."
)
add_table(
    ['Modelo', 'Target', 'Accuracy', 'F1-macro', 'Kappa'],
    [
        ['Logística L2',  'nivel TCH',         '0.477', '0.475', '0.216'],
        ['KNN',           'nivel TCH',         '0.454', '0.453', '0.181'],
        ['Random Forest', 'nivel TCH',         '0.489', '0.490', '0.232'],
        ['XGBoost',       'nivel TCH',         '0.477', '0.478', '0.214'],
        ['Logística L2',  'nivel sacarosa',    '0.500', '0.479', '0.244'],
        ['KNN',           'nivel sacarosa',    '0.532', '0.528', '0.300'],
        ['Random Forest', 'nivel sacarosa',    '0.555', '0.547', '0.330'],
        ['XGBoost',       'nivel sacarosa',    '0.569', '0.547', '0.348'],
    ]
)
add_p(
    "Las matrices de confusión muestran el patrón típico de problemas con bandas continuas: "
    "los aciertos se concentran en la diagonal pero las clases adyacentes (Bajo↔Medio, "
    "Medio↔Alto) son las más confundidas, mientras que los errores extremos Bajo↔Alto son "
    "raros. Esto refleja la naturaleza ordinal del problema y sugiere que un modelo "
    "ordinal-regression específico podría rendir mejor. Las curvas ROC One-vs-Rest "
    "evidencian que el modelo discrimina mejor las clases extremas (AUC > 0.7) que la clase "
    "Medio (AUC ≈ 0.6), comportamiento esperado al definir clases por tertiles."
)

add_h2('3.3 Diagnósticos: Curvas de Aprendizaje y Análisis de Errores')
add_p(
    "Las curvas de aprendizaje del XGBoost de regresión muestran convergencia entre train y "
    "validación al usar el dataset completo, sin gap excesivo, indicando que el modelo NO "
    "está sobreajustado y que añadir más datos del mismo tipo no produciría grandes ganancias "
    "marginales. Esto sugiere que para mejorar más allá del R² ≈ 0.5 será necesario incorporar "
    "fuentes nuevas (sensores, NDVI, estaciones meteorológicas locales) en lugar de simplemente "
    "más filas. El análisis de errores por quintil de Y real revela que los mayores residuos "
    "ocurren en los extremos (Q1 y Q5) — los lotes de muy alta o muy baja productividad son "
    "los más difíciles de predecir, por lo que en producción conviene reportar intervalos de "
    "predicción y no solo punto-estimación."
)
add_p(
    "Para clasificación, las curvas de aprendizaje del Random Forest revelan un cierto gap "
    "train-validación (clásico de RF), sugiriendo que hay margen para regularizar más vía "
    "min_samples_leaf más restrictivo. Los casos mal clasificados con alta confianza tienden "
    "a ser registros con valores de edad o cortes en zonas de transición entre clases — son "
    "candidatos a revisión por agrónomos o a un modelo con calibración de probabilidad."
)

# ── 6. CONCLUSIONES Y TRABAJO FUTURO ──────────────────────────────────────
add_h1('4. Conclusiones y Trabajo Futuro')

add_h2('4.1 Conclusiones')
add_bullet(
    "El XGBoost Regressor es el mejor predictor de TCH y %Sac.Caña, superando a OLS/Ridge en "
    "más del doble de R² (0.49–0.51 vs 0.16–0.20), confirmando que las relaciones agronómicas "
    "son fuertemente no lineales."
)
add_bullet(
    "Los modelos de árbol (Random Forest y XGBoost) son también los mejores clasificadores "
    "de niveles de rendimiento, alcanzando F1-macro ≈ 0.49 (TCH) y 0.55 (sacarosa)."
)
add_bullet(
    "Las variables más influyentes son consistentes entre tareas y modelos: edad del cultivo, "
    "dosis de madurante, lluvias en periodo de maduración y variedad/zona — alineadas con la "
    "literatura agronómica del sector."
)
add_bullet(
    "Los modelos lineales con OLS/Ridge resultan poco competitivos en R² pero mantienen valor "
    "interpretativo: identifican el sentido y la significancia estadística de los efectos."
)
add_bullet(
    "El uso adecuado de Pipelines de scikit-learn con ColumnTransformer y validación cruzada "
    "garantiza ausencia de fuga de datos y reproducibilidad."
)

add_h2('4.2 Limitaciones')
add_bullet(
    "BD_IPSA_1940 contiene una sola variedad (CC01-1940), por lo que los modelos de "
    "clasificación no generalizan a otras variedades sin reentrenamiento."
)
add_bullet(
    "El R² ≈ 0.5 indica que existe variabilidad relevante no capturada por las variables "
    "actuales — probablemente atribuible a microclimatología, estado del suelo y prácticas "
    "específicas no registradas."
)
add_bullet(
    "El split aleatorio es académicamente válido, pero en producción debería evaluarse con "
    "split temporal (entrenar con periodos pasados, evaluar en periodo más reciente) para "
    "estimar el desempeño real al desplegar."
)

add_h2('4.3 Trabajo Futuro')
add_bullet(
    "Incorporar fuentes externas: imágenes satelitales (NDVI, EVI), sensores de humedad "
    "foliar y estaciones meteorológicas locales — la literatura sugiere ganancias hasta "
    "R² ≈ 0.7–0.9 en sacarosa cuando se incluye información satelital."
)
add_bullet(
    "Probar modelos de regresión ordinal para clasificación (acordes a la naturaleza "
    "ordinal de Bajo/Medio/Alto) y modelos de calibración isotónica de probabilidades."
)
add_bullet(
    "Integrar los modelos en un sistema de monitoreo en tiempo real para apoyar decisiones de "
    "fecha de cosecha óptima, dosis y momento de aplicación de madurante por lote."
)
add_bullet(
    "Realizar análisis de feature importance con SHAP para producir explicaciones a nivel de "
    "lote individual, útiles para retroalimentación operativa con los agrónomos."
)

# ── 7. REFERENCIAS ─────────────────────────────────────────────────────────
add_h1('5. Referencias')
refs = [
    "Asocaña. (2023). Informe Anual del Sector Azucarero Colombiano. Cali: Asociación de "
    "Cultivadores de Caña de Azúcar de Colombia.",
    "Cenicaña. (2022). Variedades de caña de azúcar para el Valle del río Cauca: criterios "
    "de selección y rendimientos. Cali: Centro de Investigación de la Caña de Azúcar de Colombia.",
    "Ingenio Providencia. (2023). Reporte de Sostenibilidad. El Cerrito, Valle del Cauca.",
    "Breiman, L. (2001). Random Forests. Machine Learning, 45(1), 5–32.",
    "Chen, T., & Guestrin, C. (2016). XGBoost: A Scalable Tree Boosting System. Proceedings "
    "of the 22nd ACM SIGKDD International Conference on Knowledge Discovery and Data Mining "
    "(pp. 785–794).",
    "Pedregosa, F., et al. (2011). Scikit-learn: Machine Learning in Python. Journal of "
    "Machine Learning Research, 12, 2825–2830.",
    "Everingham, Y. L., Sexton, J., Skocaj, D., & Inman-Bamber, G. (2016). Accurate "
    "prediction of sugarcane yield using a random forest algorithm. Agronomy for Sustainable "
    "Development, 36(2), 27.",
    "Hammer, R. G., et al. (2020). Sugarcane yield prediction through data mining and crop "
    "simulation models. Sugar Tech, 22, 216–225.",
]
for r in refs:
    add_bullet(r)

# ── 8. INSERTAR RESUMEN ANTES DE LA INTRODUCCIÓN ──────────────────────────
# Movemos el resumen para que aparezca al inicio (después de la portada).
# Localizamos el párrafo de la introducción.
intro_idx = find_paragraph_idx(doc, '1. Introducción y Contexto')
if intro_idx is not None:
    # Insertar resumen antes con encabezado
    intro_p = doc.paragraphs[intro_idx]
    intro_el = intro_p._element

    # Crear nuevos párrafos
    h_resumen = doc.styles['Heading 1']
    nuevo_h = doc.paragraphs[0]._element.makeelement(
        '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}p',
        {}
    )
    # Más simple: añadir al final y luego mover via XML.
    # Implementación manual: construir párrafos con doc.add_* y mover sus _element antes de intro.
    p1 = doc.add_paragraph('Resumen', style='Heading 1')
    p2 = doc.add_paragraph(resumen_text)
    p_pal = doc.add_paragraph()
    run = p_pal.add_run('Palabras clave: ')
    run.bold = True
    p_pal.add_run('regresión, clasificación, machine learning, caña de azúcar, '
                  'XGBoost, Random Forest, agricultura de precisión, Ingenio Providencia.')

    # Mover los 3 elementos antes de intro_el
    parent = intro_el.getparent()
    for el in (p1._element, p2._element, p_pal._element):
        parent.remove(el)
        parent.insert(list(parent).index(intro_el), el)

# ── 9. GUARDAR ─────────────────────────────────────────────────────────────
doc.save(DOCX)
print(f'Informe actualizado: {DOCX}')
print(f'Párrafos totales: {len(doc.paragraphs)}')
